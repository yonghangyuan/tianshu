"""代码助手 Skill — 代码生成、文件编辑、项目分析。

三爻分类: 人（分析决策）
"""

from __future__ import annotations

from pathlib import Path

from .base import BaseSkill, SkillTool


class CodeAssistSkill(BaseSkill):
    name = "code-assist"
    description = "代码生成、文件读写、项目结构分析"
    trigram = "人"
    trigger_keywords = [
        "代码", "code", "写一个", "帮我写", "修改文件",
        "创建文件", "项目结构", "python", "Python",
    ]

    def get_tools(self) -> list[SkillTool]:
        return [
            SkillTool(
                name="write_file",
                description="Write text/Markdown content to a file. Overwrites existing files.",
                parameters={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Absolute file path"},
                        "content": {"type": "string", "description": "Content to write"},
                    },
                    "required": ["path", "content"],
                },
                handler=self._write_file,
            ),
            SkillTool(
                name="write_docx",
                description="Generate a .docx Word document from Markdown content. Supports headings, paragraphs, bold, italic, and lists.",
                parameters={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Output .docx file path"},
                        "title": {"type": "string", "description": "Document title (appears as heading)"},
                        "content": {"type": "string", "description": "Body content in Markdown format. # Heading, **bold**, *italic*, - bullet list, 1. numbered list, --- horizontal rule"},
                    },
                    "required": ["path", "title", "content"],
                },
                handler=self._write_docx,
            ),
            SkillTool(
                name="list_project",
                description="列出项目目录结构（忽略 __pycache__、.git 等）。",
                parameters={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "项目根目录路径", "default": "."},
                    },
                    "required": [],
                },
                handler=self._list_project,
            ),
        ]

    # ── 工具实现 ─────────────────────────────────────────────────────

    async def _write_file(self, path: str, content: str) -> str:
        """写入文件。"""
        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return f"已写入: {p} ({len(content)} 字符)"

    async def _write_docx(self, path: str, title: str, content: str) -> str:
        """Generate a .docx file from Markdown content."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse Markdown content
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # Horizontal rule
            if stripped == "---" or stripped == "***":
                doc.add_paragraph("─" * 60)
                continue

            # Headings
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)

            # Bullet list
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")

            # Numbered list
            elif stripped[0].isdigit() and ". " in stripped[:4]:
                doc.add_paragraph(stripped.split(". ", 1)[1], style="List Number")

            # Regular paragraph with inline formatting
            else:
                p = doc.add_paragraph()
                self._add_formatted_run(p, stripped)

        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return f"已生成: {p} ({p.stat().st_size // 1024} KB)"

    @staticmethod
    def _add_formatted_run(paragraph, text: str) -> None:
        """Parse inline **bold** and *italic* in text."""
        import re
        from docx.shared import Pt

        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    async def _list_project(self, path: str = ".") -> str:
        """列出项目结构。"""
        p = Path(path)
        if not p.exists():
            return f"路径不存在: {path}"

        skip = {"__pycache__", ".git", ".obsidian", "node_modules", ".venv", "venv"}
        lines = [f"{p}/"]

        def _walk(d: Path, indent: int) -> None:
            items = sorted(d.iterdir(), key=lambda x: (not x.is_dir(), x.name))
            for item in items:
                if item.name in skip or item.name.startswith("."):
                    continue
                prefix = "  " * indent + ("├─ " if item.is_dir() else "📄 ")
                lines.append(f"{prefix}{item.name}")
                if item.is_dir():
                    _walk(item, indent + 1)

        _walk(p, 1)
        return "\n".join(lines[:60])  # 最多 60 行
